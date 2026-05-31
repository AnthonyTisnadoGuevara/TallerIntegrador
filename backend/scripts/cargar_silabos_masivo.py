import re
import sys
import unicodedata
import uuid
from pathlib import Path

from docx import Document


DRY_RUN = False
FORZAR_NOMBRE_ARCHIVO = True

BACKEND_DIR = Path(__file__).resolve().parents[1]
UPLOADS_DIR = BACKEND_DIR / "uploads" / "silabos"
BUCKET_NAME = "silabos"

if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

from app.services.supabase_client import supabase  # noqa: E402


PALABRAS_CODIGO_SOSPECHOSO = [
    "biblioteca",
    "editorial",
    "isbn",
    "edicion",
    "edición",
    "referencia",
    "http",
    "www",
]

FRASES_ASIGNATURA_SOSPECHOSA = [
    "la asignatura",
    "de “",
    'de "',
    "es de caracter",
    "es de carácter",
    "pertenece al area",
    "pertenece al área",
    "biblioteca",
    "sumilla",
    "semestre academico",
    "semestre académico",
]

REEMPLAZOS_ASIGNATURA = {
    "BIGDATA": "BIG DATA",
    "BIG DATA": "BIG DATA",
    "DEEPLEARNIG": "DEEP LEARNING",
    "DEEPLEARNING": "DEEP LEARNING",
}


def quitar_tildes(valor: str) -> str:
    normalizado = unicodedata.normalize("NFD", valor or "")
    return "".join(caracter for caracter in normalizado if unicodedata.category(caracter) != "Mn")


def normalizar_texto(texto: str) -> str:
    return re.sub(r"\s+", " ", texto or "").strip()


def limpiar_asignatura(valor) -> str:
    nombre = Path(str(valor)).stem
    nombre = nombre.replace("_", " ").replace("-", " ")
    nombre = re.sub(r"\([^)]*\)", " ", nombre)
    nombre = re.sub(r"\b20\d{4}\b", " ", nombre)
    nombre = re.sub(r"\b\d{10,}\b", " ", nombre)
    nombre = re.sub(r"\b(ISIA|ICSI|M)\b", " ", nombre, flags=re.IGNORECASE)
    nombre = re.sub(r"\b\d{3,5}\b", " ", nombre)
    nombre = re.sub(r"\s+", " ", nombre).strip(" .-_")
    nombre = nombre.upper()
    return REEMPLAZOS_ASIGNATURA.get(nombre, nombre)


def obtener_datos_desde_nombre_archivo(nombre_archivo) -> dict:
    ruta = Path(nombre_archivo)
    stem = ruta.stem
    ciclo = None
    cuerpo = stem

    match = re.match(r"^(\d{1,2})[_-]+(.+)$", stem)
    if match:
        ciclo = int(match.group(1))
        cuerpo = match.group(2)

    asignatura = limpiar_asignatura(cuerpo)
    return {"ciclo": ciclo, "asignatura": asignatura}


def extraer_texto_docx(ruta: Path) -> str:
    document = Document(str(ruta))
    textos = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            textos.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            celdas = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if celdas:
                textos.append(" | ".join(celdas))

    return "\n".join(textos)


def limpiar_valor_extraido(valor: str) -> str:
    valor = normalizar_texto(valor)
    valor = re.split(r"\s{2,}| \| ", valor)[0]
    return valor.strip(" :-|")


def buscar_por_patrones(texto: str, patrones: list[str]) -> str | None:
    lineas = [normalizar_texto(linea) for linea in texto.splitlines() if normalizar_texto(linea)]

    for linea in lineas:
        for patron in patrones:
            match = re.search(patron, linea, flags=re.IGNORECASE)
            if match:
                valor = limpiar_valor_extraido(match.group(1))
                if valor:
                    return valor

    return None


def buscar_asignatura_segura(texto: str) -> str | None:
    patrones = [
        r"^(?:3|3\.)\s*ASIGNATURA\s*[:\-]?\s*(.+)$",
        r"^ASIGNATURA\s*[:\-]\s*(.+)$",
        r"^ASIGNATURA\s+(.+)$",
    ]
    valor = buscar_por_patrones(texto, patrones)
    if not valor:
        return None

    asignatura = limpiar_asignatura(valor)
    return asignatura if asignatura_es_valida(asignatura) and not asignatura_es_sospechosa(valor) else None


def buscar_codigo_seguro(texto: str) -> str | None:
    patrones = [
        r"^(?:5|5\.)\s*C[ÓO]DIGO\s*[:\-]?\s*(.+)$",
        r"^C[ÓO]DIGO\s*[:\-]\s*(.+)$",
        r"^C[ÓO]DIGO\s+(.+)$",
    ]
    valor = buscar_por_patrones(texto, patrones)
    if not valor:
        return None

    codigo = limpiar_codigo(valor)
    return codigo if codigo_es_valido(codigo) else None


def buscar_entero(texto: str, patrones: list[str]) -> int | None:
    valor = buscar_por_patrones(texto, patrones)
    if not valor:
        return None
    match = re.search(r"\d+", valor)
    return int(match.group(0)) if match else None


def buscar_ciclo_seguro(texto: str) -> int | None:
    ciclo = buscar_entero(
        texto,
        [
            r"^(?:6|6\.)\s*CICLO(?:\s+DE\s+ESTUDIOS)?\s*[:\-]?\s*(.+)$",
            r"^CICLO(?:\s+DE\s+ESTUDIOS)?\s*[:\-]\s*(.+)$",
        ],
    )
    return ciclo if ciclo_es_valido(ciclo) else None


def buscar_valor_general(texto: str, etiqueta: str) -> str | None:
    patron = rf"^{re.escape(etiqueta)}\s*[:\-]\s*(.+)$"
    return buscar_por_patrones(texto, [patron])


def buscar_fecha(texto: str, etiqueta: str) -> str | None:
    valor = buscar_valor_general(texto, etiqueta)
    if not valor:
        return None

    match = re.search(r"(\d{4})[-/](\d{1,2})[-/](\d{1,2})", valor)
    if match:
        anio, mes, dia = match.groups()
        return f"{anio}-{int(mes):02d}-{int(dia):02d}"

    match = re.search(r"(\d{1,2})[-/](\d{1,2})[-/](\d{4})", valor)
    if match:
        dia, mes, anio = match.groups()
        return f"{anio}-{int(mes):02d}-{int(dia):02d}"

    return None


def buscar_correo(texto: str) -> str | None:
    correos = re.findall(r"[\w.\-+]+@[\w.\-]+\.\w+", texto)
    return ", ".join(dict.fromkeys(correos)) if correos else None


def limpiar_codigo(valor: str) -> str:
    codigo = normalizar_texto(valor)
    codigo = re.split(r"\s{2,}| \| |;", codigo)[0]
    codigo = codigo.strip(" :-|.,")
    return codigo.upper()


def codigo_es_valido(codigo: str | None) -> bool:
    if not codigo:
        return False
    if len(codigo) > 30:
        return False
    normalizado = quitar_tildes(codigo).lower()
    if any(palabra in normalizado for palabra in [quitar_tildes(p).lower() for p in PALABRAS_CODIGO_SOSPECHOSO]):
        return False
    if re.search(r"\s{2,}", codigo):
        return False
    return bool(re.search(r"[A-Z]", codigo) and re.search(r"\d", codigo))


def asignatura_es_sospechosa(valor: str | None) -> bool:
    normalizado = quitar_tildes(valor or "").lower()
    return any(quitar_tildes(frase).lower() in normalizado for frase in FRASES_ASIGNATURA_SOSPECHOSA)


def asignatura_es_valida(asignatura: str | None) -> bool:
    if not asignatura:
        return False
    if len(asignatura) > 120:
        return False
    if asignatura_es_sospechosa(asignatura):
        return False
    return bool(re.search(r"[A-ZÁÉÍÓÚÜÑ]", asignatura))


def ciclo_es_valido(ciclo) -> bool:
    return isinstance(ciclo, int) and 1 <= ciclo <= 10


def generar_codigo_asignatura(ciclo: int, asignatura: str) -> str:
    palabras = re.findall(r"[A-ZÁÉÍÓÚÜÑ0-9]+", asignatura.upper())
    iniciales = "".join(palabra[0] for palabra in palabras[:4]) or "SIL"
    return f"C{ciclo:02d}-{iniciales}"


def normalizar_codigo_final(codigo: str | None, ciclo: int, asignatura: str) -> str:
    if codigo_es_valido(codigo):
        return codigo
    generado = generar_codigo_asignatura(ciclo, asignatura)
    return generado if codigo_es_valido(generado) else "SIN-CODIGO"


def extraer_metadatos_docx(ruta: Path) -> dict:
    texto = extraer_texto_docx(ruta)
    return {
        "texto_extraido": texto,
        "semestre_academico": buscar_valor_general(texto, "SEMESTRE ACADÉMICO")
        or buscar_valor_general(texto, "SEMESTRE ACADEMICO"),
        "facultad": buscar_valor_general(texto, "FACULTAD"),
        "programa_estudios": buscar_valor_general(texto, "PROGRAMA DE ESTUDIOS"),
        "asignatura": buscar_asignatura_segura(texto),
        "codigo_asignatura": buscar_codigo_seguro(texto),
        "ciclo": buscar_ciclo_seguro(texto),
        "modalidad": buscar_valor_general(texto, "MODALIDAD"),
        "creditos": buscar_entero(texto, [r"^(?:7|7\.)\s*CR[ÉE]DITOS\s*[:\-]?\s*(.+)$", r"^CR[ÉE]DITOS\s*[:\-]\s*(.+)$"]),
        "total_horas_semestrales": buscar_entero(texto, [r"^TOTAL\s+DE\s+HORAS\s+SEMESTRALES\s*[:\-]\s*(.+)$"]),
        "total_horas_semanales": buscar_entero(texto, [r"^TOTAL\s+DE\s+HORAS\s+SEMANALES\s*[:\-]\s*(.+)$"]),
        "fecha_inicio": buscar_fecha(texto, "FECHA DE INICIO"),
        "fecha_culminacion": buscar_fecha(texto, "FECHA DE CULMINACIÓN") or buscar_fecha(texto, "FECHA DE CULMINACION"),
        "duracion_semanas": buscar_entero(texto, [r"^DURACI[ÓO]N(?:\s+SEMANAS)?\s*[:\-]\s*(.+)$"]),
        "docente_responsable": buscar_valor_general(texto, "DOCENTE RESPONSABLE") or buscar_valor_general(texto, "DOCENTE"),
        "correo_docente": buscar_correo(texto),
    }


def extraer_metadatos_archivo(ruta: Path) -> dict:
    datos_nombre = obtener_datos_desde_nombre_archivo(ruta.name)
    metadatos = {}

    if ruta.suffix.lower() == ".docx":
        try:
            metadatos = extraer_metadatos_docx(ruta)
        except Exception as error:
            print(f"  - No se pudo extraer metadatos DOCX de {ruta.name}: {error}")
    elif ruta.suffix.lower() == ".pdf":
        print(f"  - PDF detectado sin extractor configurado, se usaran datos del nombre: {ruta.name}")

    ciclo = datos_nombre["ciclo"] if FORZAR_NOMBRE_ARCHIVO and datos_nombre["ciclo"] else metadatos.get("ciclo")
    ciclo = ciclo or datos_nombre["ciclo"]

    asignatura_documento = metadatos.get("asignatura")
    asignatura_nombre = datos_nombre["asignatura"]
    asignatura = asignatura_nombre if FORZAR_NOMBRE_ARCHIVO and asignatura_nombre else asignatura_documento
    asignatura = asignatura if asignatura_es_valida(asignatura) else asignatura_nombre

    if not ciclo_es_valido(ciclo):
        ciclo = metadatos.get("ciclo")

    if not ciclo_es_valido(ciclo):
        raise ValueError("No se pudo determinar un ciclo válido entre 1 y 10.")
    if not asignatura_es_valida(asignatura):
        raise ValueError("No se pudo determinar una asignatura válida.")

    codigo = normalizar_codigo_final(metadatos.get("codigo_asignatura"), ciclo, asignatura)

    return {
        "semestre_academico": metadatos.get("semestre_academico") or "202510",
        "facultad": metadatos.get("facultad") or "Ingeniería",
        "programa_estudios": metadatos.get("programa_estudios") or "Ingeniería de Sistemas e Inteligencia Artificial",
        "asignatura": asignatura[:120],
        "codigo_asignatura": codigo[:30],
        "ciclo": ciclo,
        "modalidad": metadatos.get("modalidad") or "Presencial",
        "creditos": metadatos.get("creditos"),
        "total_horas_semestrales": metadatos.get("total_horas_semestrales"),
        "total_horas_semanales": metadatos.get("total_horas_semanales"),
        "fecha_inicio": metadatos.get("fecha_inicio"),
        "fecha_culminacion": metadatos.get("fecha_culminacion"),
        "duracion_semanas": metadatos.get("duracion_semanas"),
        "docente_responsable": metadatos.get("docente_responsable"),
        "correo_docente": metadatos.get("correo_docente"),
    }


def obtener_content_type(ruta: Path) -> str:
    if ruta.suffix.lower() == ".docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if ruta.suffix.lower() == ".pdf":
        return "application/pdf"
    return "application/octet-stream"


def subir_archivo_storage(ruta: Path) -> str:
    nombre_storage = f"{uuid.uuid4()}_{ruta.name}"
    ruta_storage = f"silabos/{nombre_storage}"

    with ruta.open("rb") as archivo:
        supabase.storage.from_(BUCKET_NAME).upload(
            path=ruta_storage,
            file=archivo.read(),
            file_options={
                "content-type": obtener_content_type(ruta),
                "upsert": "true",
            },
        )

    return supabase.storage.from_(BUCKET_NAME).get_public_url(ruta_storage)


def buscar_silabo_existente(datos: dict) -> dict | None:
    response = (
        supabase.table("silabos")
        .select("*")
        .eq("asignatura", datos["asignatura"])
        .eq("ciclo", datos["ciclo"])
        .limit(1)
        .execute()
    )
    return response.data[0] if response.data else None


def construir_registro(ruta: Path) -> dict:
    datos = extraer_metadatos_archivo(ruta)
    archivo_url = f"DRY_RUN://{ruta.name}" if DRY_RUN else subir_archivo_storage(ruta)

    return {
        **datos,
        "archivo_url": archivo_url,
        "estado": "pendiente",
        "porcentaje_cumplimiento": 0,
        "observacion_general": "Registro cargado masivamente desde carpeta local.",
    }


def validar_registro(datos: dict) -> None:
    if not ciclo_es_valido(datos.get("ciclo")):
        raise ValueError("El ciclo debe ser entero entre 1 y 10.")
    if not asignatura_es_valida(datos.get("asignatura")):
        raise ValueError("La asignatura está vacía, es sospechosa o supera 120 caracteres.")
    if not codigo_es_valido(datos.get("codigo_asignatura")) and datos.get("codigo_asignatura") != "SIN-CODIGO":
        raise ValueError("El código de asignatura es sospechoso.")


def procesar_archivo(ruta: Path) -> str:
    datos = construir_registro(ruta)
    validar_registro(datos)
    existente = buscar_silabo_existente(datos)
    accion = "actualizado" if existente else "insertado"

    print(
        f"Archivo: {ruta.name} | ciclo={datos['ciclo']} | "
        f"asignatura={datos['asignatura']} | codigo={datos['codigo_asignatura']} | "
        f"accion={accion}"
    )

    if DRY_RUN:
        return f"dry_run_{accion}"

    if existente:
        supabase.table("silabos").update(datos).eq("id", existente["id"]).execute()
        return "actualizado"

    supabase.table("silabos").insert(datos).execute()
    return "insertado"


def listar_archivos() -> list[Path]:
    if not UPLOADS_DIR.exists():
        raise FileNotFoundError(f"No existe la carpeta de carga: {UPLOADS_DIR}")

    return sorted(
        [
            ruta
            for ruta in UPLOADS_DIR.iterdir()
            if ruta.is_file() and ruta.suffix.lower() in {".docx", ".pdf"}
        ],
        key=lambda ruta: ruta.name.lower(),
    )


def main() -> None:
    archivos = listar_archivos()
    validos = 0
    omitidos = 0
    insertados = 0
    actualizados = 0
    errores = []

    print(f"DRY_RUN: {DRY_RUN}")
    print(f"FORZAR_NOMBRE_ARCHIVO: {FORZAR_NOMBRE_ARCHIVO}")
    print(f"Archivos encontrados: {len(archivos)}")

    for ruta in archivos:
        try:
            resultado = procesar_archivo(ruta)
            validos += 1
            if resultado.endswith("insertado"):
                insertados += 1
            elif resultado.endswith("actualizado"):
                actualizados += 1
        except ValueError as error:
            omitidos += 1
            errores.append((ruta.name, str(error)))
            print(f"OMITIDO {ruta.name}: {error}")
        except Exception as error:
            errores.append((ruta.name, str(error)))
            print(f"ERROR {ruta.name}: {error}")

    print("\nResumen de carga masiva")
    print(f"Archivos encontrados: {len(archivos)}")
    print(f"Válidos: {validos}")
    print(f"Omitidos: {omitidos}")
    print(f"Insertados: {insertados}")
    print(f"Actualizados: {actualizados}")
    print(f"Errores: {len(errores)}")

    if errores:
        print("\nDetalle de errores/omitidos:")
        for nombre, error in errores:
            print(f"- {nombre}: {error}")


if __name__ == "__main__":
    main()
