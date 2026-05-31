import os
import re
import tempfile
import unicodedata
from io import BytesIO
from urllib.parse import parse_qs, quote, urlparse
from zipfile import BadZipFile

import requests
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from pypdf import PdfReader


def convertir_google_drive_url(url: str) -> str:
    """Convierte enlaces compartidos de Google Drive a enlaces de descarga."""
    if not url or not isinstance(url, str):
        return url

    parsed_url = urlparse(url.strip())
    if parsed_url.netloc.lower() not in {"drive.google.com", "www.drive.google.com"}:
        return url.strip()

    match = re.search(r"/file/d/([^/]+)", parsed_url.path)
    file_id = match.group(1) if match else parse_qs(parsed_url.query).get("id", [None])[0]

    if not file_id:
        return url.strip()

    return f"https://drive.google.com/uc?export=download&id={quote(file_id, safe='')}"


def descargar_archivo_con_metadata(url: str) -> tuple[bytes, str, str]:
    """Descarga un archivo documental desde Drive, Storage o una URL directa."""
    url_descarga = convertir_google_drive_url(url)
    es_google_drive = "drive.google.com" in urlparse(url_descarga).netloc.lower()

    try:
        response = requests.get(url_descarga, timeout=30)
        response.raise_for_status()
    except requests.RequestException as error:
        if es_google_drive:
            raise RuntimeError(
                "No se pudo descargar el archivo de Google Drive. "
                'Verifique que esté compartido como "Cualquier persona con el enlace puede ver".'
            ) from error
        raise RuntimeError(f"No se pudo descargar el archivo desde la URL indicada: {error}") from error

    content_type = response.headers.get("content-type", "").lower()
    contenido_inicial = response.content[:200].lstrip().lower()
    es_html = contenido_inicial.startswith((b"<!doctype html", b"<html"))
    if "text/html" in content_type or es_html:
        if es_google_drive:
            raise RuntimeError(
                "Google Drive no entregó el documento. "
                'Verifique que esté compartido como "Cualquier persona con el enlace puede ver".'
            )
        raise RuntimeError("La URL indicada devolvió una página HTML en lugar de un archivo documental.")

    return response.content, content_type, url_descarga


def descargar_archivo_desde_url(url: str) -> bytes:
    """Descarga un archivo documental desde Drive, Storage o una URL directa."""
    contenido, _, _ = descargar_archivo_con_metadata(url)
    return contenido


def extraer_texto_docx(contenido: bytes) -> str:
    """Extrae texto de un archivo DOCX desde bytes."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        temp_file.write(contenido)
        temp_path = temp_file.name

    try:
        try:
            document = Document(temp_path)
        except (BadZipFile, PackageNotFoundError, ValueError, KeyError) as error:
            raise ValueError(
                "El archivo descargado no es un documento .docx válido o no puede leerse."
            ) from error

        textos = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                textos.append(paragraph.text.strip())

        for table in document.tables:
            for row in table.rows:
                celdas = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if celdas:
                    textos.append(" | ".join(celdas))

        return "\n".join(textos)
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def extraer_texto_docx_desde_url(url: str) -> str:
    """Descarga un archivo DOCX público y extrae texto de párrafos y tablas."""
    contenido = descargar_archivo_desde_url(url)
    return extraer_texto_docx(contenido)


def extraer_texto_pdf(contenido_bytes: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(contenido_bytes))
        textos = []

        for page in reader.pages:
            texto = page.extract_text()
            if texto:
                textos.append(texto)

        texto_final = "\n".join(textos).strip()

        if not texto_final:
            raise ValueError(
                "El PDF no contiene texto extraíble. Puede ser un documento escaneado. "
                "Para analizarlo, conviértalo a DOCX o use un PDF con texto seleccionable."
            )

        return texto_final
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"No se pudo extraer texto del PDF: {str(e)}") from e


def _es_pdf(url: str, content_type: str, contenido: bytes) -> bool:
    ruta = urlparse(url.strip()).path.lower()
    return (
        ruta.endswith(".pdf")
        or ".pdf" in url.lower()
        or "application/pdf" in (content_type or "").lower()
        or contenido[:5] == b"%PDF-"
    )


def extraer_texto_desde_url(url: str) -> str:
    """Extrae texto desde un documento remoto DOCX o PDF."""
    if not url or not isinstance(url, str):
        raise ValueError("No se proporcionó una URL de documento válida.")

    url_limpia = url.strip()
    contenido, content_type, url_descarga = descargar_archivo_con_metadata(url_limpia)
    ruta = urlparse(url_descarga).path.lower()
    extension = os.path.splitext(ruta)[1]

    if _es_pdf(url_limpia, content_type, contenido):
        return extraer_texto_pdf(contenido)

    if extension == ".doc":
        raise ValueError(
            "El formato .doc antiguo no está soportado. Convierta el archivo a .docx o PDF con texto seleccionable."
        )

    if extension and extension != ".docx" and "/file/d/" not in ruta:
        raise ValueError("Solo se permite validar documentos .docx o PDF con texto extraíble.")

    texto = extraer_texto_docx(contenido)
    if not texto.strip():
        raise ValueError("El documento .docx no contiene texto extraíble.")

    return texto


def validar_secciones_silabo(texto: str) -> list[dict]:
    """Valida si el texto del sílabo contiene las secciones obligatorias."""
    texto_normalizado = unicodedata.normalize("NFD", texto.lower())
    texto_normalizado = "".join(
        caracter for caracter in texto_normalizado if unicodedata.category(caracter) != "Mn"
    )

    secciones = [
        {
            "seccion": "Datos generales",
            "palabras_clave": ["datos generales", "facultad", "programa", "asignatura", "codigo"],
        },
        {"seccion": "Sumilla", "palabras_clave": ["sumilla"]},
        {
            "seccion": "Aporte al perfil de egreso",
            "palabras_clave": ["aporte de la asignatura", "perfil de egreso", "competencia"],
        },
        {
            "seccion": "Programación por unidades de aprendizaje",
            "palabras_clave": ["programacion por unidades", "unidad", "semanas", "contenidos"],
        },
        {
            "seccion": "Estrategias metodológicas",
            "palabras_clave": ["estrategias metodologicas"],
        },
        {
            "seccion": "Recursos y escenarios de enseñanza-aprendizaje",
            "palabras_clave": ["recursos", "escenarios", "ensenanza", "aprendizaje"],
        },
        {
            "seccion": "Técnicas e instrumentos de evaluación",
            "palabras_clave": ["tecnicas", "instrumentos de evaluacion", "evaluacion"],
        },
        {
            "seccion": "Estrategias de tutoría y apoyo pedagógico",
            "palabras_clave": ["tutoria", "apoyo pedagogico"],
        },
        {"seccion": "Bibliografía", "palabras_clave": ["bibliografia"]},
    ]

    resultados = []
    for item in secciones:
        cumple = any(palabra in texto_normalizado for palabra in item["palabras_clave"])
        resultados.append(
            {
                "seccion": item["seccion"],
                "cumple": cumple,
                "observacion": (
                    "Sección identificada en el documento."
                    if cumple
                    else "No se identificó la sección en el documento."
                ),
            }
        )

    return resultados
